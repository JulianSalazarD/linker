"""
Extractor para archivos .docx (python-docx).
"""
from __future__ import annotations

import io
from pathlib import Path

from PIL import Image

from processor.extractor import BaseExtractor, SourceContent, register_extractor
from processor.ocr import ocr_image


class DocxExtractor(BaseExtractor):
    supported = (".docx",)

    def extract(self, source: str | Path, *, ocr: bool = False) -> SourceContent:
        from docx import Document

        doc = Document(str(source))

        paragraphs = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        text = "\n".join(p for p in paragraphs if p.strip())

        images: list[Image.Image] = []
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    images.append(Image.open(io.BytesIO(rel.target_part.blob)))
                except Exception:
                    pass

        if ocr and images:
            ocr_texts = [ocr_image(img) for img in images]
            text = "\n".join(filter(None, [text] + ocr_texts))

        return SourceContent(
            text=text,
            images=images,
            source=str(source),
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )


register_extractor(DocxExtractor())

import math
import os
import platform
import subprocess
from docxtpl import DocxTemplate, RichText, InlineImage
from docx.shared import Mm
from io import BytesIO
from datetime import date
import re
import tempfile
import json
from pathlib import Path
from docx import Document

_MESES = [
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]

def _fmt(n) -> str:
    """Formatea un número como entero con separador de miles (.)."""
    return f"{int(n):,}".replace(",", ".")


def convert_to_pdf(docx_path: Path, prefer_libreoffice: bool = False, prefer_onlyoffice: bool = False) -> Path | None:
    """Convierte un DOCX a PDF usando docx2pdf (Windows), OnlyOffice o LibreOffice.
    Retorna el Path del PDF o None si falla."""
    docx_path = Path(docx_path)

    # docx2pdf solo funciona en Windows
    if not prefer_libreoffice and not prefer_onlyoffice and platform.system() == "Windows":
        try:
            from docx2pdf import convert
            convert(str(docx_path), str(docx_path.with_suffix(".pdf")))
            pdf_path = docx_path.with_suffix(".pdf")
            return pdf_path if pdf_path.exists() else None
        except Exception:
            pass

    # Intentar OnlyOffice Document Server CLI
    if prefer_onlyoffice:
        for cmd in ["documentserver-convert", "/usr/bin/documentserver-convert"]:
            try:
                subprocess.run(
                    [cmd, "--input", str(docx_path), "--output", str(docx_path.with_suffix(".pdf"))],
                    check=True,
                    capture_output=True,
                    timeout=60,
                )
                pdf_path = docx_path.with_suffix(".pdf")
                if pdf_path.exists():
                    return pdf_path
            except Exception:
                pass

    # Intentar LibreOffice (Linux/macOS/Windows con LibreOffice instalado)
    for cmd in ["soffice", "libreoffice", "/usr/bin/soffice"]:
        try:
            subprocess.run(
                [cmd, "--headless", "--convert-to", "pdf", "--outdir", str(docx_path.parent), str(docx_path)],
                check=True,
                capture_output=True,
                timeout=60,
            )
            pdf_path = docx_path.with_suffix(".pdf")
            if pdf_path.exists():
                return pdf_path
        except Exception:
            pass

    return None


_IF_BLOCK = re.compile(r'\{%-?\s*if\s+\w+\s*-?%\}.*\{%-?\s*endif\s*-?%\}', re.DOTALL)
_SENTINEL = "\u00abCOND\u00bb"  # «COND» — marca párrafos condicionales

CONFIG_PATH = Path("config/data.json")


if platform.system() == "Windows":
    import msvcrt

    def _lock(f):
        msvcrt.locking(f.fileno(), msvcrt.LK_LOCK, 1)

    def _unlock(f):
        try:
            f.seek(0)
            msvcrt.locking(f.fileno(), msvcrt.LK_UNLCK, 1)
        except OSError:
            pass
else:
    import fcntl

    def _lock(f):
        fcntl.flock(f.fileno(), fcntl.LOCK_EX)

    def _unlock(f):
        fcntl.flock(f.fileno(), fcntl.LOCK_UN)


def _increment_quote_number() -> str:
    """Atomically reads and increments the quote number in config/data.json.
    Uses file locking to prevent race conditions."""
    lock_path = CONFIG_PATH.with_suffix(".lock")
    with open(lock_path, "w") as lockf:
        _lock(lockf)
        try:
            with open(CONFIG_PATH, "r") as f:
                config = json.load(f)
            current = int(config.get("number", "0"))
            new_number = str(current + 1).zfill(5)
            config["number"] = new_number
            tmp_path = CONFIG_PATH.with_suffix(".tmp")
            with open(tmp_path, "w") as f:
                json.dump(config, f, indent=2)
            Path(tmp_path).replace(CONFIG_PATH)
            return new_number
        finally:
            _unlock(lockf)


def _prepare_template(template_path: str) -> str:
    """Copia el template a un temporal añadiendo el sentinel al inicio
    de cada párrafo condicional {% if %}...{% endif %}.
    Retorna la ruta del archivo temporal."""
    doc = Document(str(template_path))
    for para in doc.paragraphs:
        if _IF_BLOCK.search(para.text) and para.runs:
            para.runs[0].text = _SENTINEL + para.runs[0].text
    fd, tmp = tempfile.mkstemp(suffix=".docx")
    import os; os.close(fd)
    doc.save(tmp)
    return tmp


def _remove_null_paragraphs(output_path: Path):
    """Elimina párrafos donde el sentinel quedó solo (condición False)
    y limpia el sentinel de los párrafos donde la condición fue True."""
    doc = Document(str(output_path))
    to_remove = []
    for para in doc.paragraphs:
        if para.text.strip() == _SENTINEL:
            to_remove.append(para)
        elif _SENTINEL in para.text:
            for run in para.runs:
                if _SENTINEL in run.text:
                    run.text = run.text.replace(_SENTINEL, "")
                    break
    for para in to_remove:
        para._element.getparent().remove(para._element)
    doc.save(str(output_path))


def fill_template(
    data: dict,
    products: list,
    template_path: str = "config/template_base.docx",
    output_dir: str = "pruebas",
    output_name: str = "DOCUMENTO PROCEDIMIENTO INSTALACIÓN",
    fotos: list[dict] | None = None,
    generate_pdf: bool = False,
    prefer_libreoffice: bool = False,
    prefer_onlyoffice: bool = False,
) -> tuple[Path, Path | None]:
    """Rellena template_base.docx con los datos extraídos y guarda el resultado.

    Returns:
        tuple[Path, Path | None]: (path_al_documento_docx, path_al_pdf_o_None)
    """
    tmp_template = _prepare_template(template_path)

    try:
        tpl = DocxTemplate(tmp_template)

        # Convertir imágenes PIL a InlineImage de docxtpl
        if fotos:
            for foto in fotos:
                buf = BytesIO()
                foto["imagen"].save(buf, format="PNG")
                buf.seek(0)
                foto["imagen"] = InlineImage(tpl, image_descriptor=buf, width=Mm(150))

        hoy = date.today()
        fecha = f"{hoy.day} de {_MESES[hoy.month - 1]} de {hoy.year}"

        ciudad_dep = data.get("ciudad_departtamento") or data.get("ciudad_departamento") or "Bogotá D.C."
        ciudad = ciudad_dep.split(",")[0].strip()

        total_iva = sum(p.iva for p in products)
        valor_total = sum(p.vr_total for p in products)
        valor_inicial = math.floor(valor_total / 2 / 100_000) * 100_000
        number = _increment_quote_number()

        context = {
            "NOMBRE":              (data.get("nombre") or "").upper(),
            "is_male":             "Señor" if data.get("is_male", True) else "Señora",
            "direccion":           data.get("direccion", ""),
            "apartamento":         data.get("apartamento"),
            "nombre_edificio":     data.get("nombre_edificio"),
            "ciudad_departamento": ciudad_dep,
            "ciudad":              ciudad,
            "fecha":               fecha,
            "date":                hoy.strftime("%Y-%m-%d"),
            "tipo_vehiculo":       data.get("tipo_vehiculo") or "Tesla",
            "fotos":               fotos or [],
            "new_page":            RichText("\f"),
            "number":              number,
            "products":            [
                type("P", (), dict(
                    item=p.item, qty=p.qty, desc=p.desc, unit=p.unit,
                    value=_fmt(p.value), iva=_fmt(p.iva), vr_total=_fmt(p.vr_total)
                ))() for p in products
            ],
            "total_iva":           _fmt(total_iva),
            "valor_total":         _fmt(valor_total),
            "valor_inicial":        f"{valor_inicial:,}",
        }

        tpl.render(context)

        output_path = Path(output_dir) / f"{output_name}.docx"
        tpl.save(output_path)
    finally:
        Path(tmp_template).unlink(missing_ok=True)

    _remove_null_paragraphs(output_path)

    pdf_path = None
    if generate_pdf:
        pdf_path = convert_to_pdf(output_path, prefer_libreoffice=prefer_libreoffice, prefer_onlyoffice=prefer_onlyoffice)
        if pdf_path is None:
            print("[red]Error generando PDF[/]")

    return output_path, pdf_path